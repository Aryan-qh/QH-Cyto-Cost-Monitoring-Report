import requests
from datetime import datetime, timedelta
from tabulate import tabulate
import time
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load environment variables
load_dotenv()

# Azure AD Configuration
tenant_id = os.getenv('AZURE_TENANT_ID')
client_id = os.getenv('AZURE_CLIENT_ID')
client_secret = os.getenv('AZURE_CLIENT_SECRET')

# Subscription Configuration
subscriptions = {
    'main': os.getenv('SUBSCRIPTION_MAIN'),
    'prod': os.getenv('SUBSCRIPTION_PROD'),
    'dev': os.getenv('SUBSCRIPTION_DEV'),
    'test': os.getenv('SUBSCRIPTION_TEST')
}

# Validate required environment variables
required_vars = ['AZURE_TENANT_ID', 'AZURE_CLIENT_ID', 'AZURE_CLIENT_SECRET', 
                 'SUBSCRIPTION_MAIN', 'SUBSCRIPTION_PROD', 'SUBSCRIPTION_DEV', 'SUBSCRIPTION_TEST']
missing_vars = [var for var in required_vars if not os.getenv(var)]
if missing_vars:
    raise ValueError(f"Missing required environment variables: {', '.join(missing_vars)}")

# Authenticate with Azure AD and get access token
auth_url = f'https://login.microsoftonline.com/{tenant_id}/oauth2/token'
auth_data = {
    'grant_type': 'client_credentials',
    'client_id': client_id,
    'client_secret': client_secret,
    'resource': 'https://management.azure.com/'
}
auth_response = requests.post(auth_url, data=auth_data)
access_token = auth_response.json()['access_token']

def get_cost_data(subscription_id, date, retry_count=0, max_retries=3):
    """Get cost data for a specific subscription and date with retry logic"""
    usage_url = f'https://management.azure.com/subscriptions/{subscription_id}/providers/Microsoft.CostManagement/query?api-version=2023-03-01'
    
    usage_data = {
        'type': 'Usage',
        'timeframe': 'Custom',
        'timePeriod': {
            'from': date.strftime('%Y-%m-%dT00:00:00Z'),
            'to': date.strftime('%Y-%m-%dT23:59:59Z')
        },
        'dataset': {
            'granularity': 'Daily',
            'aggregation': {
                'totalCost': {
                    'name': 'Cost',
                    'function': 'Sum'
                }
            },
            'grouping': [
                {
                    'type': 'Dimension',
                    'name': 'ResourceType'
                }
            ]
        }
    }
    
    try:
        response = requests.post(
            usage_url, 
            headers={'Authorization': f'Bearer {access_token}'}, 
            json=usage_data
        )
        
        # Handle rate limiting
        if response.status_code == 429:
            if retry_count < max_retries:
                # Get retry-after header or use exponential backoff
                retry_after = int(response.headers.get('Retry-After', 2 ** retry_count))
                print(f"Rate limit hit. Waiting {retry_after} seconds before retry...")
                time.sleep(retry_after)
                return get_cost_data(subscription_id, date, retry_count + 1, max_retries)
            else:
                print(f"Max retries reached for {date.strftime('%Y-%m-%d')}")
                return []
        
        response.raise_for_status()
        return response.json()['properties']['rows']
    except Exception as e:
        print(f"Error fetching data for {date.strftime('%Y-%m-%d')}: {str(e)}")
        return []

def get_cost_data_range(subscription_id, start_date, end_date):
    """Get cost data for a date range in a single API call"""
    usage_url = f'https://management.azure.com/subscriptions/{subscription_id}/providers/Microsoft.CostManagement/query?api-version=2023-03-01'
    
    usage_data = {
        'type': 'Usage',
        'timeframe': 'Custom',
        'timePeriod': {
            'from': start_date.strftime('%Y-%m-%dT00:00:00Z'),
            'to': end_date.strftime('%Y-%m-%dT23:59:59Z')
        },
        'dataset': {
            'granularity': 'Daily',
            'aggregation': {
                'totalCost': {
                    'name': 'Cost',
                    'function': 'Sum'
                }
            },
            'grouping': [
                {
                    'type': 'Dimension',
                    'name': 'ResourceType'
                },
                {
                    'type': 'Dimension',
                    'name': 'ChargeType'
                }
            ]
        }
    }
    
    try:
        response = requests.post(
            usage_url, 
            headers={'Authorization': f'Bearer {access_token}'}, 
            json=usage_data,
            timeout=30
        )
        
        # Handle rate limiting
        if response.status_code == 429:
            retry_after = int(response.headers.get('Retry-After', 60))
            print(f"Rate limit hit. Waiting {retry_after} seconds...")
            time.sleep(retry_after)
            return get_cost_data_range(subscription_id, start_date, end_date)
        
        response.raise_for_status()
        return response.json()['properties']
    except Exception as e:
        print(f"Error fetching data range: {str(e)}")
        return None

def process_cost_data(raw_data):
    """Process raw cost data into categories based on Resource Type"""
    costs = {
        'Databricks': 0,
        'Virtual Machine': 0,
        'Storage': 0,
        'Others': 0
    }
    
    for row in raw_data:
        cost = row[0]
        resource_type = row[2].lower() if len(row) > 2 else ''
        
        # Categorize based on resource type
        if 'databricks/workspaces' in resource_type or 'databricks/workspace' in resource_type:
            costs['Databricks'] += cost
        elif 'compute/virtualmachines' in resource_type or 'microsoft.compute/virtualmachines' in resource_type:
            costs['Virtual Machine'] += cost
        elif 'storage/storageaccounts' in resource_type or 'microsoft.storage/storageaccounts' in resource_type:
            costs['Storage'] += cost
        else:
            costs['Others'] += cost
    
    return costs

def parse_range_response(response_data, num_days):
    """Parse the range API response and organize by date"""
    if not response_data or 'rows' not in response_data:
        return {}
    
    # Get column indices
    columns = response_data.get('columns', [])
    cost_idx = next((i for i, col in enumerate(columns) if col['name'] == 'Cost'), 0)
    date_idx = next((i for i, col in enumerate(columns) if col['name'] == 'UsageDate'), 1)
    resource_idx = next((i for i, col in enumerate(columns) if col['name'] == 'ResourceType'), 2)
    
    # Organize data by date
    daily_data = {}
    for row in response_data['rows']:
        date = row[date_idx]
        if date not in daily_data:
            daily_data[date] = []
        daily_data[date].append(row)
    
    return daily_data

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'tc{}'.format(edge.capitalize())
            element = OxmlElement('w:{}'.format(tag))
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))
            tcPr.append(element)

def add_table_to_doc(doc, table_data, headers, title=None):
    """Add a formatted table to the Word document"""
    if title:
        para = doc.add_paragraph()
        run = para.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add data rows
    for row_data in table_data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Add spacing

def generate_data_for_subscription(subscription_id, subscription_name, num_days):
    """Generate cost data for a subscription and return it"""
    print(f"\nFetching data for {subscription_name.upper()} subscription...")
    
    # Calculate date range
    end_date = datetime.now() - timedelta(days=1)
    start_date = end_date - timedelta(days=num_days - 1)
    
    # Get all data in one API call
    response_data = get_cost_data_range(subscription_id, start_date, end_date)
    
    if not response_data:
        print(f"Failed to fetch data for {subscription_name} subscription.\n")
        return None, None
    
    daily_data = parse_range_response(response_data, num_days)
    
    # Prepare table data
    cost_table_data = []
    percent_table_data = []
    all_costs = []
    date_strings = []
    
    # Process each day
    for i in range(num_days - 1, -1, -1):
        date = datetime.now() - timedelta(days=i+1)
        date_key = int(date.strftime('%Y%m%d'))
        date_str = date.strftime('%m/%d')
        date_strings.append(date_str)
        
        # Get data for this date
        day_rows = daily_data.get(date_key, [])
        costs = process_cost_data(day_rows)
        all_costs.append(costs)
    
    # Determine which categories have data (skip Databricks for main)
    categories = ['Databricks', 'Virtual Machine', 'Storage', 'Others']
    if subscription_name.lower() == 'main':
        # Check if main has any Databricks costs
        has_databricks = any(costs['Databricks'] > 0 for costs in all_costs)
        if not has_databricks:
            categories = ['Virtual Machine', 'Storage', 'Others']
    
    # Build cost table
    for i, costs in enumerate(all_costs):
        row = [date_strings[i]]
        for category in categories:
            row.append(f"${costs[category]:.2f}")
        cost_table_data.append(row)
    
    # Build percentage change table
    for i in range(1, len(all_costs)):
        row = [date_strings[i]]
        
        for category in categories:
            prev_cost = all_costs[i-1][category]
            curr_cost = all_costs[i][category]
            
            if prev_cost == 0:
                if curr_cost == 0:
                    percent_change = 0
                else:
                    percent_change = 100
            else:
                percent_change = ((curr_cost - prev_cost) / prev_cost) * 100
            
            row.append(f"{percent_change:+.2f}%")
        
        percent_table_data.append(row)
    
    headers = ['Date'] + categories
    
    return {
        'cost_table': cost_table_data,
        'percent_table': percent_table_data,
        'headers': headers,
        'date_strings': date_strings
    }

def create_word_document(all_data, num_days):
    """Create a Word document with all the cost data"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Azure Cost Summary Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create date range string for greeting
    end_date = datetime.now() - timedelta(days=1)
    start_date = end_date - timedelta(days=num_days - 1)
    
    # Format dates with day names
    date_list = []
    for i in range(num_days):
        date = start_date + timedelta(days=i)
        date_list.append(f"{date.strftime('%A')} ({date.strftime('%m/%d')})")
    
    date_range_str = ", ".join(date_list)
    
    # Add greeting
    greeting = doc.add_paragraph()
    greeting.add_run("Hi Team,\n\n").bold = False
    greeting.add_run(f"Please find below the Azure cost summary for {date_range_str} for all subscriptions, along with percentage changes compared to the previous day.\n")
    
    # Add tables for each subscription in order: prod, dev, test, main
    for sub_name in ['prod', 'dev', 'test', 'main']:
        if sub_name in all_data and all_data[sub_name]:
            data = all_data[sub_name]
            
            # Add subscription header
            sub_header = doc.add_heading(f'{sub_name.capitalize()} Environment', level=2)
            
            # Add cost table
            add_table_to_doc(doc, data['cost_table'], data['headers'])
            
            # Add percentage difference table
            add_table_to_doc(doc, data['percent_table'], data['headers'], 
                           f"Percentage difference for {sub_name}")
    
    # Add closing
    doc.add_paragraph("\nThank you.")
    
    # Save document
    filename = f"Azure_Cost_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    print(f"\n{'='*80}")
    print(f"Word document created: {filename}")
    print(f"{'='*80}")
    
    return filename

def generate_table_for_subscription(subscription_id, subscription_name, num_days):
    """Generate cost table for a subscription (console output)"""
    print(f"\n{'='*80}")
    print(f"{subscription_name.upper()} SUBSCRIPTION")
    print(f"{'='*80}\n")
    
    # Calculate date range
    end_date = datetime.now() - timedelta(days=1)
    start_date = end_date - timedelta(days=num_days - 1)
    
    print(f"Fetching data from {start_date.strftime('%Y-%m-%d')} to {end_date.strftime('%Y-%m-%d')}...")
    
    # Get all data in one API call
    response_data = get_cost_data_range(subscription_id, start_date, end_date)
    
    if not response_data:
        print("Failed to fetch data for this subscription.\n")
        return
    
    daily_data = parse_range_response(response_data, num_days)
    
    # Prepare table data
    table_data = []
    all_costs = []
    
    # Process each day
    for i in range(num_days - 1, -1, -1):
        date = datetime.now() - timedelta(days=i+1)
        date_key = int(date.strftime('%Y%m%d'))
        
        # Get data for this date
        day_rows = daily_data.get(date_key, [])
        costs = process_cost_data(day_rows)
        all_costs.append(costs)
        
        # Create row
        row = [
            date.strftime('%m/%d'),
            f"${costs['Databricks']:.2f}",
            f"${costs['Virtual Machine']:.2f}",
            f"${costs['Storage']:.2f}",
            f"${costs['Others']:.2f}"
        ]
        
        table_data.append(row)
    
    # Define headers
    headers = ['Date', 'Databricks', 'Virtual Machine', 'Storage', 'Others']
    categories = ['Databricks', 'Virtual Machine', 'Storage', 'Others']
    
    # Print cost table
    print("Cost Table:")
    print(tabulate(table_data, headers=headers, tablefmt='grid'))
    print()
    
    # Calculate and print percentage change table
    print("Percentage Change (Day over Day):")
    percent_table_data = []
    
    for i in range(1, len(all_costs)):
        date = (datetime.now() - timedelta(days=len(all_costs)-i)).strftime('%m/%d')
        row = [date]
        
        for category in categories:
            prev_cost = all_costs[i-1][category]
            curr_cost = all_costs[i][category]
            
            if prev_cost == 0:
                if curr_cost == 0:
                    percent_change = 0
                else:
                    percent_change = 100
            else:
                percent_change = ((curr_cost - prev_cost) / prev_cost) * 100
            
            row.append(f"{percent_change:+.2f}%")
        
        percent_table_data.append(row)
    
    print(tabulate(percent_table_data, headers=headers, tablefmt='grid'))
    print()

if __name__ == "__main__":
    # Get user input for number of days
    while True:
        try:
            num_days = int(input("Enter the number of days to look back (ending at yesterday): "))
            if num_days < 1:
                print("Please enter a positive number.")
                continue
            if num_days > 90:
                print("Warning: Requesting more than 90 days may take a long time.")
                confirm = input("Do you want to continue? (yes/no): ")
                if confirm.lower() != 'yes':
                    continue
            break
        except ValueError:
            print("Please enter a valid number.")

    # Generate tables for all subscriptions (console output)
    print("\n" + "="*80)
    print(f"AZURE COST REPORT - LAST {num_days} DAYS (ending yesterday)")
    print("="*80)

    # Collect data for Word document
    all_data = {}
    
    # Process subscriptions with delay between them
    for idx, sub_name in enumerate(['main', 'prod', 'dev', 'test']):
        # Generate console output
        generate_table_for_subscription(subscriptions[sub_name], sub_name.upper(), num_days)
        
        # Collect data for Word document
        data = generate_data_for_subscription(subscriptions[sub_name], sub_name, num_days)
        if data:
            all_data[sub_name] = data
        
        # Add delay between subscriptions to avoid rate limiting (except for the last one)
        if idx < 3:
            print("Waiting 2 seconds before next subscription...")
            time.sleep(2)

    print("="*80)
    print("Console report generation completed!")
    print("="*80)
    
    # Create Word document
    print("\nGenerating Word document...")
    create_word_document(all_data, num_days)
    
    print("\nAll tasks completed successfully!")